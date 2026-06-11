"""doc-security-mapper: document content extraction script.

Supported formats: PDF, DOCX, DOC, XLSX, XLS, CSV, MD, TXT.
Usage: python extract.py <file_path>
"""

import csv
import os
import sys


def extract_pdf(path):
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            text = ""
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            if text.strip():
                return text
    except ImportError:
        pass

    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    except Exception as e:
        raise RuntimeError(f"PDF 解析失败: {e}")


def extract_docx(path):
    try:
        from docx import Document

        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"DOCX 解析失败: {e}")


def extract_xlsx(path):
    try:
        import openpyxl

        wb = openpyxl.load_workbook(path, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    parts.append(row_text)

        return "\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"XLSX 解析失败: {e}")


def extract_csv(path):
    parts = []
    encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312"]
    for enc in encodings:
        try:
            with open(path, encoding=enc, newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    parts.append(" | ".join(row))
            return "\n".join(parts)
        except UnicodeDecodeError:
            continue

    raise RuntimeError("CSV 编码解析失败")


def extract_text(path):
    encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue

    raise RuntimeError("文本文件编码解析失败")


def extract(path):
    if not os.path.exists(path):
        raise FileNotFoundError(f"文件不存在: {path}")

    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return extract_pdf(path)
    if ext in (".docx",):
        return extract_docx(path)
    if ext in (".doc",):
        raise RuntimeError("旧版 .doc 格式需先转换为 .docx，抱歉，无法解析")
    if ext in (".xlsx", ".xlsm"):
        return extract_xlsx(path)
    if ext in (".xls",):
        raise RuntimeError("旧版 .xls 格式需先转换为 .xlsx，抱歉，无法解析")
    if ext == ".csv":
        return extract_csv(path)
    if ext in (".md", ".txt", ".text", ".log", ".rst"):
        return extract_text(path)

    raise RuntimeError(f"不支持的文件格式: {ext}，抱歉，无法解析")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python extract.py <文件路径>", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]
    try:
        content = extract(file_path)
        if not content or not content.strip():
            print("抱歉，无法解析该文档（文档内容为空或无法提取文本）", file=sys.stderr)
            sys.exit(1)
        print(content)
    except Exception as e:
        print(f"抱歉，无法解析该文档: {e}", file=sys.stderr)
        sys.exit(1)
